from __future__ import annotations

import logging
import os
import re
import tempfile
from typing import Optional

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import JSONResponse

logger = logging.getLogger("studybridge.extract_compat")
router = APIRouter()


def _error_json(status_code: int, error_code: str, message: str, recovery_suggestion: str) -> JSONResponse:
    """스펙 계약(error_code/recoverable/recovery_suggestion) 구조화 에러 응답.

    상태코드는 비2xx를 유지해(Spring의 기존 에러 분기 호환) 본문에만 error_code 계약을 싣는다.
    """
    return JSONResponse(
        status_code=status_code,
        content={
            "status": "ERROR",
            "extraction_status": "FAILED",
            "error_code": error_code,
            "message": message,
            "recoverable": True,
            "recovery_suggestion": recovery_suggestion,
        },
    )

TEXT_ENCODINGS = ("utf-8", "cp949", "euc-kr", "latin-1")
PDF_TEXT_MIN_CHARS = int(os.getenv("PDF_MIN_TEXT_CHARS", "80"))


def _extract_with_pymupdf(path: str) -> tuple[str, int]:
    import fitz

    parts: list[str] = []

    with fitz.open(path) as doc:
        page_count = len(doc)
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            text = text.strip()
            if text:
                parts.append(f"\n\n[Page {index}]\n{text}")

    return "\n".join(parts).strip(), page_count


def _extract_with_pypdf(path: str) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    parts: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(f"\n\n[Page {index}]\n{text}")

    return "\n".join(parts).strip(), len(reader.pages)


def _extract_pdf_text(path: str) -> tuple[str, int, str]:
    best_text = ""
    best_count = 0
    best_engine = "none"

    try:
        text, page_count = _extract_with_pymupdf(path)
        best_text, best_count, best_engine = text, page_count, "pymupdf"
        if len(text.strip()) >= PDF_TEXT_MIN_CHARS:
            return text, page_count, "pymupdf"
        logger.warning("PyMuPDF extraction short chars=%s threshold=%s; trying pypdf/OCR", len(text.strip()), PDF_TEXT_MIN_CHARS)
    except Exception as first_error:
        logger.warning("PyMuPDF extraction failed: %s", first_error)

    try:
        text, page_count = _extract_with_pypdf(path)
        if len(text.strip()) > len(best_text.strip()):
            best_text, best_count, best_engine = text, page_count, "pypdf"
        if len(text.strip()) >= PDF_TEXT_MIN_CHARS:
            return text, page_count, "pypdf"
        logger.warning("pypdf extraction short chars=%s threshold=%s; trying OCR", len(text.strip()), PDF_TEXT_MIN_CHARS)
    except Exception as second_error:
        logger.warning("pypdf extraction failed: %s", second_error)

    try:
        from app.services.ocr_service import extract_pdf_ocr_from_path, should_attempt_ocr
        if should_attempt_ocr(best_text, PDF_TEXT_MIN_CHARS):
            ocr = extract_pdf_ocr_from_path(path, min_chars=PDF_TEXT_MIN_CHARS)
            if ocr.text and len(ocr.text.strip()) > len(best_text.strip()):
                return ocr.text, ocr.page_count or best_count, ocr.engine
            logger.warning("OCR fallback did not improve extraction engine=%s reason=%s", ocr.engine, ocr.reason)
    except Exception as ocr_error:
        logger.warning("OCR fallback exception: %s", ocr_error)

    if best_text.strip():
        return best_text, best_count, best_engine

    raise HTTPException(
        status_code=422,
        detail="No extractable text found in PDF; OCR unavailable or produced too little text",
    )


def _extract_docx_text(path: str) -> tuple[str, int, str]:
    try:
        from docx import Document
    except Exception as e:
        logger.exception("python-docx is not available")
        raise HTTPException(status_code=500, detail=f"DOCX extraction dependency failed: {type(e).__name__}")

    doc = Document(path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                rows.append(row_text)
        if rows:
            parts.append(f"[Table {table_index}]\n" + "\n".join(rows))

    return "\n\n".join(parts).strip(), len(parts), "python-docx"


def _extract_txt_text(path: str) -> tuple[str, int, str]:
    last_error: Exception | None = None
    for encoding in TEXT_ENCODINGS:
        try:
            with open(path, "r", encoding=encoding) as f:
                text = f.read()
            section_count = len([p for p in re.split(r"\n\s*\n", text) if p.strip()]) or 1
            return text.strip(), section_count, f"txt-{encoding}"
        except UnicodeDecodeError as e:
            last_error = e

    logger.warning("TXT decoding failed for every fallback encoding: %s", last_error)
    raise HTTPException(status_code=500, detail="TXT extraction failed: unsupported encoding")


def _detect_file_type(filename: str, content_type: Optional[str]) -> str:
    lower = filename.lower()
    content_type = (content_type or "").lower()
    if lower.endswith(".pdf") or content_type == "application/pdf":
        return "pdf"
    if lower.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    if lower.endswith(".txt") or content_type.startswith("text/plain"):
        return "txt"
    # 구형 .doc(OLE2)는 미지원 — 스펙상 별도 error_code로 안내한다.
    if lower.endswith(".doc") or content_type == "application/msword":
        return "doc"
    return "unsupported"


@router.post("/api/extract")
async def extract_document_compat(
    file: UploadFile = File(...),
    material_id: Optional[int] = Form(None),
):
    filename = os.path.basename(file.filename or "uploaded")
    file_type = _detect_file_type(filename, file.content_type)

    if file_type == "doc":
        logger.warning("[DOCUMENT EXTRACT UNSUPPORTED] filename=%s reason=legacy_doc", filename)
        return _error_json(
            400,
            "UNSUPPORTED_DOC_FORMAT",
            "현재는 .docx 형식만 지원합니다. .doc 파일은 .docx로 변환 후 업로드해주세요.",
            "MS Word 또는 한컴에서 다른 이름으로 저장을 사용해 .docx로 변환하세요.",
        )
    if file_type == "unsupported":
        logger.warning("[DOCUMENT EXTRACT UNSUPPORTED] filename=%s reason=unsupported_type", filename)
        return _error_json(
            400,
            "UNSUPPORTED_FILE_TYPE",
            "지원하지 않는 파일 형식입니다. PDF 또는 DOCX 파일을 업로드해주세요.",
            "PDF 또는 DOCX 형식으로 변환 후 다시 업로드하세요.",
        )

    suffix = os.path.splitext(filename)[1].lower() or f".{file_type}"
    temp_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            temp_path = tmp.name
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                tmp.write(chunk)

        if file_type == "pdf":
            text, count, engine = _extract_pdf_text(temp_path)
            count_field = "page_count"
        elif file_type == "docx":
            text, count, engine = _extract_docx_text(temp_path)
            count_field = "section_count"
        else:
            text, count, engine = _extract_txt_text(temp_path)
            count_field = "section_count"

        if not text.strip():
            logger.warning(
                "[DOCUMENT EXTRACT EMPTY] material_id=%s filename=%s type=%s reason=no_text",
                material_id,
                filename,
                file_type,
            )
            if file_type == "docx":
                return _error_json(
                    422,
                    "DOCX_TEXT_EMPTY",
                    "DOCX 문서에서 분석 가능한 텍스트를 찾지 못했습니다.",
                    "문서 본문에 텍스트가 포함되어 있는지 확인하거나 PDF/DOCX 원본을 다시 업로드하세요.",
                )
            raise HTTPException(status_code=422, detail=f"No extractable text found in {file_type.upper()}")

        logger.info(
            "[DOCUMENT EXTRACT OK] material_id=%s filename=%s type=%s chars=%s engine=%s",
            material_id,
            filename,
            file_type,
            len(text),
            engine,
        )

        response = {
            "status": "SUCCESS",
            "extraction_status": "SUCCESS",
            "filename": filename,
            "material_id": material_id,
            "file_type": file_type,
            "engine": engine,
            "extracted_text": text,
            "extractedText": text,
            "text": text,
            "content": text,
            "page_count": count if count_field == "page_count" else None,
            "section_count": count if count_field == "section_count" else None,
        }
        return response

    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
